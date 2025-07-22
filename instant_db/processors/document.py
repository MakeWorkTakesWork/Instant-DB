"""
Document processor for Instant-DB
Handles individual document processing and metadata extraction
"""

import hashlib
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime

from ..core.database import InstantDB
from ..core.chunking import ChunkingEngine, TextChunk


class DocumentProcessor:
    """
    Process individual documents for Instant-DB
    Handles document parsing, chunking, and indexing
    """
    
    def __init__(self, 
                 embedding_provider: str = "sentence-transformers",
                 embedding_model: str = "all-MiniLM-L6-v2",
                 vector_db: str = "chroma",
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200):
        """
        Initialize document processor
        
        Args:
            embedding_provider: Embedding provider to use
            embedding_model: Embedding model name
            vector_db: Vector database type
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
        """
        self.embedding_provider = embedding_provider
        self.embedding_model = embedding_model
        self.vector_db = vector_db
        
        # Initialize chunking engine
        self.chunking_engine = ChunkingEngine(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # Supported file types
        self.supported_extensions = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.html': 'text/html',
            '.htm': 'text/html'
        }
    
    def process_document(self, 
                        file_path: Union[str, Path], 
                        output_dir: Optional[str] = None,
                        metadata: Optional[Dict] = None) -> Dict[str, Any]:
        """
        Process a single document
        
        Args:
            file_path: Path to document file
            output_dir: Output directory for database
            metadata: Additional metadata to attach
            
        Returns:
            Processing results dictionary
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "status": "error",
                "error": f"File not found: {file_path}",
                "file_path": str(file_path)
            }
        
        if not self._is_supported_file(file_path):
            return {
                "status": "error", 
                "error": f"Unsupported file type: {file_path.suffix}",
                "file_path": str(file_path),
                "supported_types": list(self.supported_extensions.keys())
            }
        
        try:
            # Extract text from document
            text_content = self._extract_text(file_path)
            
            if not text_content or not text_content.strip():
                return {
                    "status": "error",
                    "error": "No text content extracted from document",
                    "file_path": str(file_path)
                }
            
            # Generate document metadata
            doc_metadata = self._generate_metadata(file_path, text_content, metadata)
            
            # Create document ID
            document_id = self._generate_document_id(file_path, text_content)
            
            # Chunk the document
            chunks = self.chunking_engine.chunk_text(
                text=text_content,
                document_id=document_id,
                metadata=doc_metadata
            )
            
            if not chunks:
                return {
                    "status": "error",
                    "error": "No chunks generated from document",
                    "file_path": str(file_path)
                }
            
            # Initialize database
            db_path = output_dir or "./instant_db_database"
            db = InstantDB(
                db_path=db_path,
                embedding_provider=self.embedding_provider,
                embedding_model=self.embedding_model,
                vector_db=self.vector_db
            )
            
            # Convert chunks to document format for database
            documents = []
            for chunk in chunks:
                doc = {
                    "id": chunk.id,
                    "content": chunk.content,
                    "document_id": document_id,
                    "chunk_index": chunk.chunk_index,
                    "section": chunk.section,
                    "subsection": chunk.subsection,
                    "chunk_type": chunk.chunk_type,
                    "word_count": chunk.word_count,
                    "char_count": chunk.char_count,
                    **chunk.metadata
                }
                documents.append(doc)
            
            # Add to database
            db_result = db.add_documents(documents)
            
            # Get chunking statistics
            chunk_stats = self.chunking_engine.get_chunking_stats(chunks)
            
            return {
                "status": "success",
                "file_path": str(file_path),
                "document_id": document_id,
                "database_path": db_path,
                "chunks_processed": len(chunks),
                "total_chars": len(text_content),
                "total_words": len(text_content.split()),
                "chunk_stats": chunk_stats,
                "db_result": db_result,
                "metadata": doc_metadata
            }
            
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file_path": str(file_path),
                "exception_type": type(e).__name__
            }
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file type is supported"""
        return file_path.suffix.lower() in self.supported_extensions
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension in ['.txt', '.md']:
                return self._extract_text_plain(file_path)
            elif file_extension == '.pdf':
                return self._extract_text_pdf(file_path)
            elif file_extension in ['.docx']:
                return self._extract_text_docx(file_path)
            elif file_extension in ['.html', '.htm']:
                return self._extract_text_html(file_path)
            else:
                # Fallback to plain text
                return self._extract_text_plain(file_path)
                
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from {file_path}: {e}")
    
    def _extract_text_plain(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _extract_text_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
        except ImportError:
            try:
                import pypdf2 as PyPDF2
            except ImportError:
                raise ImportError(
                    "PDF processing requires PyPDF2 or pypdf2. "
                    "Install with: pip install PyPDF2"
                )
        
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
        
        return text.strip()
    
    def _extract_text_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "DOCX processing requires python-docx. "
                "Install with: pip install python-docx"
            )
        
        doc = docx.Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text += " | ".join(row_text) + "\n"
        
        return text.strip()
    
    def _extract_text_html(self, file_path: Path) -> str:
        """Extract text from HTML files"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "HTML processing requires BeautifulSoup4. "
                "Install with: pip install beautifulsoup4"
            )
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Break into lines and remove leading/trailing space
        lines = (line.strip() for line in text.splitlines())
        
        # Break multi-headlines into lines
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        
        # Drop blank lines
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _generate_metadata(self, file_path: Path, content: str, 
                          user_metadata: Optional[Dict] = None) -> Dict[str, Any]:
        """Generate metadata for document"""
        metadata = {
            "source_file": str(file_path),
            "filename": file_path.name,
            "file_extension": file_path.suffix,
            "file_size": file_path.stat().st_size,
            "processed_at": datetime.now().isoformat(),
            "content_hash": hashlib.md5(content.encode()).hexdigest(),
            "char_count": len(content),
            "word_count": len(content.split()),
            "line_count": len(content.splitlines())
        }
        
        # Add MIME type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            metadata["mime_type"] = mime_type
        
        # Add document type classification
        metadata["document_type"] = self._classify_document_type(file_path, content)
        
        # Add user-provided metadata
        if user_metadata:
            metadata.update(user_metadata)
        
        return metadata
    
    def _classify_document_type(self, file_path: Path, content: str) -> str:
        """Classify document type based on content and filename"""
        filename_lower = file_path.name.lower()
        content_lower = content.lower()
        
        # Check for specific document types
        if any(term in filename_lower for term in ['readme', 'documentation', 'manual']):
            return "Documentation"
        elif any(term in filename_lower for term in ['meeting', 'notes', 'minutes']):
            return "Meeting Notes"  
        elif any(term in filename_lower for term in ['report', 'analysis']):
            return "Report"
        elif any(term in filename_lower for term in ['pitch', 'presentation', 'deck']):
            return "Presentation"
        elif any(term in filename_lower for term in ['email', 'message']):
            return "Email"
        elif file_path.suffix.lower() == '.pdf':
            # Look for academic paper indicators
            if any(term in content_lower for term in ['abstract', 'introduction', 'conclusion', 'references']):
                return "Research Paper"
            else:
                return "PDF Document"
        elif file_path.suffix.lower() in ['.md', '.markdown']:
            return "Markdown Document"
        else:
            return "Text Document"
    
    def _generate_document_id(self, file_path: Path, content: str) -> str:
        """Generate unique document ID"""
        # Use file path and content hash for uniqueness
        path_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        return f"doc_{path_hash}_{content_hash}"
    
    def get_supported_formats(self) -> Dict[str, str]:
        """Get supported file formats"""
        return self.supported_extensions.copy() 